"""文档上传前检查工具

用法（从项目根目录执行）：
    $env:PYTHONPATH="backend"; python scripts/check_documents.py

或在 bash：
    PYTHONPATH=backend python scripts/check_documents.py

功能：
- 弹出文件选择框
- 检查文件是否存在、大小、扩展名
- 对 .docx 检查是否是合法 ZIP、是否能被 python-docx 解析、relationships 是否异常
- 对 .pdf  检查是否能被 pymupdf 打开、各页文本量
- 对图片检查基本文件信息
- 最后弹出结果摘要
"""
import os
import sys
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox

# 把 backend 加入 Python 路径，以便复用现有解析器
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "backend"))

from app.core.config import settings


def choose_file() -> str | None:
    """弹出文件选择框，返回用户选择的文件路径。"""
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    root.attributes("-topmost", True)  # 让对话框置顶

    project_root = Path(__file__).resolve().parent.parent
    uploads_dir = project_root / "uploads"
    initial_dir = str(uploads_dir) if uploads_dir.exists() else str(project_root)

    file_path = filedialog.askopenfilename(
        title="选择要检查的文档",
        initialdir=initial_dir,
        filetypes=[
            ("文档", "*.pdf *.docx *.doc *.jpg *.jpeg *.png"),
            ("PDF", "*.pdf"),
            ("Word", "*.docx *.doc"),
            ("图片", "*.jpg *.jpeg *.png"),
            ("所有文件", "*.*"),
        ],
    )
    root.destroy()
    return file_path if file_path else None


def check_docx(file_path: str) -> list[str]:
    """检查 .docx 文件，返回问题列表（空表示无明显问题）。"""
    import zipfile
    from xml.etree import ElementTree as ET

    issues: list[str] = []

    # 1. 是否是合法 ZIP（.docx 本质是 zip）
    if not zipfile.is_zipfile(file_path):
        issues.append("该文件不是合法的 .docx ZIP 包，可能是 .doc 老格式被强制改后缀。")
        return issues

    # 2. 关键文件是否存在
    required_files = ["word/document.xml", "[Content_Types].xml"]
    try:
        with zipfile.ZipFile(file_path, "r") as zf:
            namelist = zf.namelist()
            for req in required_files:
                if req not in namelist:
                    issues.append(f"缺少关键文件：{req}，文件可能损坏。")

            # 3. 检查 relationships 里是否有 Target="NULL" 等异常
            rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
            for name in namelist:
                if name.endswith(".rels"):
                    try:
                        data = zf.read(name)
                        root = ET.fromstring(data)
                        for rel in root.findall(f"{{{rels_ns}}}Relationship"):
                            target = rel.get("Target")
                            if target is None or target.upper() == "NULL" or target == "":
                                issues.append(f"{name} 中存在非法关系 Target={target!r}（常见于 WPS）。")
                    except Exception as e:
                        issues.append(f"解析 {name} 失败：{e}")
    except zipfile.BadZipFile as e:
        issues.append(f"ZIP 包损坏：{e}")
        return issues

    # 4. 用 python-docx 实际解析
    try:
        from docx import Document

        doc = Document(file_path)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        text_preview = "\n".join(p.text.strip() for p in doc.paragraphs[:3] if p.text.strip())
        if para_count == 0 and table_count == 0:
            issues.append("文档解析成功，但没有任何段落和表格，内容为空。")
        else:
            print(f"  段落数：{para_count}")
            print(f"  表格数：{table_count}")
            if text_preview:
                print("  前 3 段预览：")
                print(text_preview[:300])
    except Exception as e:
        issues.append(f"python-docx 解析失败：{type(e).__name__}: {e}")

    return issues


def check_pdf(file_path: str) -> list[str]:
    """检查 .pdf 文件。"""
    import fitz

    issues: list[str] = []
    try:
        doc = fitz.open(file_path)
        page_count = len(doc)
        print(f"  页数：{page_count}")

        total_text = 0
        low_text_pages = []
        for i in range(page_count):
            text = doc[i].get_text().strip()
            total_text += len(text)
            if len(text) < settings.PDF_OCR_MIN_TEXT_LENGTH:
                low_text_pages.append(i + 1)

        print(f"  总文本字符数：{total_text}")
        if low_text_pages:
            issues.append(
                f"第 {low_text_pages} 页文本量过低（<{settings.PDF_OCR_MIN_TEXT_LENGTH} 字符），"
                f"这些页将走 OCR，可能耗时且识别质量取决于图片清晰度。"
            )
        if total_text == 0:
            issues.append("PDF 未提取到任何文本，可能是纯扫描件或图片 PDF。")

        doc.close()
    except Exception as e:
        issues.append(f"PDF 解析失败：{type(e).__name__}: {e}")

    return issues


def check_image(file_path: str) -> list[str]:
    """检查图片文件。"""
    issues: list[str] = []
    size = os.path.getsize(file_path)
    print(f"  文件大小：{size / 1024:.1f} KB")
    if size == 0:
        issues.append("图片文件大小为 0，文件已损坏。")
    return issues


def main():
    file_path = choose_file()
    if not file_path:
        print("未选择文件，退出。")
        return

    file_path = str(Path(file_path).resolve())
    ext = Path(file_path).suffix.lower()
    size = os.path.getsize(file_path)

    print("=" * 60)
    print(f"检查文件：{file_path}")
    print(f"扩展名：{ext}")
    print(f"大小：{size / 1024:.1f} KB")
    print("=" * 60)

    issues: list[str] = []

    if ext not in settings.ALLOWED_EXTENSIONS:
        issues.append(f"不允许的扩展名：{ext}，允许的类型：{settings.ALLOWED_EXTENSIONS}")

    if size == 0:
        issues.append("文件大小为 0。")

    if ext == ".docx":
        issues.extend(check_docx(file_path))
    elif ext == ".pdf":
        issues.extend(check_pdf(file_path))
    elif ext in {".jpg", ".jpeg", ".png"}:
        issues.extend(check_image(file_path))
    elif ext == ".doc":
        issues.append(".doc 是 Word 老二进制格式，系统暂不支持，请转换为 .docx 后上传。")
    else:
        issues.append(f"未实现 {ext} 的详细检查。")

    print("=" * 60)
    if issues:
        print("发现以下问题：")
        for idx, issue in enumerate(issues, 1):
            print(f"  {idx}. {issue}")
        summary = "发现 " + str(len(issues)) + " 个问题，详见控制台输出。"
        messagebox.showwarning("检查未通过", summary)
    else:
        print("未发现明显问题，该文件大概率可以正常上传、分块、向量化。")
        messagebox.showinfo("检查通过", "未发现明显问题，该文件可以正常上传。")


if __name__ == "__main__":
    main()
