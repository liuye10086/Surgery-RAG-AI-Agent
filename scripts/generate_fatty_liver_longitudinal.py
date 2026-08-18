from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from datetime import date, timedelta
from pathlib import Path
from typing import Any, Iterable

import numpy as np
from docx import Document


SEED = 20260818
GENERATOR_VERSION = "1.3.1"
DATA_CUTOFF = date(2026, 8, 18)
EXCLUDED_SOURCE_CASES = {"A27-1", "A30-1", "A32-1", "A34-1", "A35-1"}
PATIENT_HEADERS = [
    "patient_id", "age", "sex", "cohort_group", "fatty_liver_date",
    "final_stage", "cirrhosis_date", "hcc_date", "last_followup_date",
    "lost_to_followup",
]
VISIT_HEADERS = [
    "patient_id", "visit_date", "alt", "ast", "ggt", "tbil", "alb",
    "plt", "hba1c", "afp", "waist", "bmi",
]
INDICATORS = VISIT_HEADERS[2:]
SAFETY_BOUNDS = {
    "alt": (1.0, 2000.0), "ast": (1.0, 2000.0), "ggt": (1.0, 2500.0),
    "tbil": (0.5, 800.0), "alb": (10.0, 65.0), "plt": (10.0, 1000.0),
    "hba1c": (3.0, 20.0), "afp": (0.1, 50000.0), "waist": (45.0, 180.0),
    "bmi": (12.0, 75.0),
}
DEFAULT_BASELINES = {
    "alt": 42.0, "ast": 34.0, "ggt": 48.0, "tbil": 13.0, "alb": 44.0,
    "plt": 230.0, "hba1c": 6.0, "afp": 3.8, "waist": 91.0, "bmi": 26.8,
}
NUMBER_PATTERN = r"([0-9]{1,3}(?:\s+[0-9]{3})+(?:\.[0-9]+)?|[0-9]+(?:\.[0-9]+)?)"
LAB_PATTERNS = {
    "alt": [rf"(?:ALT|谷丙转氨酶|丙氨酸(?:氨基)?转氨酶)[：:\s]*{NUMBER_PATTERN}"],
    "ast": [rf"(?:AST|谷草转氨酶|天门冬氨酸(?:氨基)?转氨酶)[：:\s]*{NUMBER_PATTERN}"],
    "ggt": [rf"(?:GGT|γ-?谷氨酰(?:基)?转移酶|谷氨酰转肽酶)[：:\s]*{NUMBER_PATTERN}"],
    "tbil": [rf"(?:TBIL|TB|总胆红素)[：:\s]*{NUMBER_PATTERN}"],
    "alb": [rf"(?:Alb|ALB|白蛋白)[：:\s]*{NUMBER_PATTERN}"],
    "plt": [rf"(?:PLT|血小板(?:计数)?)[：:\s]*{NUMBER_PATTERN}"],
    "hba1c": [rf"(?:HbA\s*1c|HbA1c|糖化血红蛋白)[：:\s]*{NUMBER_PATTERN}\s*%?"],
    "afp": [rf"(?:AFP|甲胎蛋白)[：:\s]*{NUMBER_PATTERN}"],
    "waist": [rf"(?:腰围|腹围)[：:\s]*{NUMBER_PATTERN}\s*cm"],
    "bmi": [rf"(?:BMI|体重指数)[：:\s]*{NUMBER_PATTERN}"],
}


@dataclass(frozen=True)
class LabAnchor:
    indicator: str
    value: float
    context: str
    anchor_date: date | None = None
    sequence: int = 0


@dataclass
class CaseRecord:
    patient_id: str
    source: str
    source_number: int
    source_occurrence: int
    source_case_id: str
    paragraphs: list[str]
    full_text: str
    age: int | None
    sex: str | None
    explicit_dates: list[date] = field(default_factory=list)
    lab_anchors: list[LabAnchor] = field(default_factory=list)
    diagnosis_text: str = ""
    cohort_group: str = "mixed"
    classification_reasons: list[str] = field(default_factory=list)
    evidence_score: float = 0.0


@dataclass(frozen=True)
class GenerationConfig:
    seed: int = SEED
    min_visits: int = 3
    max_visits: int = 6
    min_span_months: int = 24
    max_span_months: int = 60
    cirrhosis_count: int = 50
    hcc_count: int = 25


def _segment_document(source: str, path: Path) -> list[dict[str, Any]]:
    paragraphs = [paragraph.text.strip() for paragraph in Document(str(path)).paragraphs]
    cases: list[dict[str, Any]] = []
    occurrences: dict[int, int] = defaultdict(int)
    current: dict[str, Any] | None = None
    for text in paragraphs:
        match = re.fullmatch(r"病例\s*(\d+)", text)
        if match:
            if current:
                cases.append(current)
            number = int(match.group(1))
            occurrences[number] += 1
            current = {
                "source": source,
                "source_number": number,
                "source_occurrence": occurrences[number],
                "paragraphs": [],
            }
        elif current is not None and text:
            current["paragraphs"].append(text)
    if current:
        cases.append(current)
    return cases


def _extract_age(paragraphs: list[str]) -> int | None:
    opening = " ".join(paragraphs[:10])
    patterns = [
        r"(?:患者|男性|女性|男|女|孕妇|产妇)[，、\s　]*(?:男|女|男性|女性)?[，、\s　]*([1-9][0-9]?)\s*岁",
        r"([1-9][0-9]?)\s*岁(?:男性|女性|男|女|患者)",
    ]
    for pattern in patterns:
        match = re.search(pattern, opening)
        if match:
            age = int(match.group(1))
            if 1 <= age <= 100:
                return age
    return None


def _extract_sex(paragraphs: list[str]) -> str | None:
    opening = " ".join(paragraphs[:8])
    female_patterns = [
        r"患者[，、\s　]*女", r"患者[，、\s　]*女性", r"女性[，、\s　]*\d+岁",
        r"\d+\s*岁\s*女性患者", r"孕妇", r"产妇",
    ]
    male_patterns = [
        r"患者[，、\s　]*男", r"患者[，、\s　]*男性", r"男性[，、\s　]*\d+岁",
        r"\d+\s*岁\s*男性患者",
    ]
    if any(re.search(pattern, opening) for pattern in female_patterns):
        return "female"
    if any(re.search(pattern, opening) for pattern in male_patterns):
        return "male"
    return None


def _safe_date(year: int, month: int = 1, day: int = 15) -> date | None:
    try:
        value = date(year, month, day)
    except ValueError:
        return None
    return value if date(1990, 1, 1) <= value <= DATA_CUTOFF else None


def _extract_dates(text: str) -> list[date]:
    values: set[date] = set()
    full_patterns = [
        r"(20\d{2})[年/\-.](\d{1,2})[月/\-.](\d{1,2})日?",
        r"(20\d{2})/(\d{1,2})/(\d{1,2})",
    ]
    for pattern in full_patterns:
        for year, month, day in re.findall(pattern, text):
            parsed = _safe_date(int(year), int(month), int(day))
            if parsed:
                values.add(parsed)
    occupied_year_months = {(item.year, item.month) for item in values}
    for year, month in re.findall(r"(20\d{2})[年/\-.](\d{1,2})月?", text):
        key = (int(year), int(month))
        if key not in occupied_year_months:
            parsed = _safe_date(key[0], key[1], 15)
            if parsed:
                values.add(parsed)
    return sorted(values)


def _date_mentions(text: str, initial_year: int | None = None) -> list[tuple[int, date]]:
    mentions: list[tuple[int, date]] = []
    current_year: int | None = initial_year
    token_pattern = re.compile(
        r"(?P<full>(?P<year>20\d{2})[年/\-.](?P<month>\d{1,2})[月/\-.](?P<day>\d{1,2})日?)"
        r"|(?P<partial>(?P<pmonth>\d{1,2})月(?P<pday>\d{1,2})(?:日|至\d{1,2}日)?)"
    )
    for match in token_pattern.finditer(text):
        if match.group("full"):
            current_year = int(match.group("year"))
            parsed = _safe_date(current_year, int(match.group("month")), int(match.group("day")))
        elif current_year is not None:
            parsed = _safe_date(current_year, int(match.group("pmonth")), int(match.group("pday")))
        else:
            parsed = None
        if parsed:
            mentions.append((match.start(), parsed))
    return mentions


def _extract_lab_anchors(paragraphs: list[str]) -> list[LabAnchor]:
    anchors: list[LabAnchor] = []
    seen: set[tuple[str, float]] = set()
    rolling_year: int | None = None
    for paragraph in paragraphs:
        normalized = paragraph.replace("：", ":").replace(" ", " ")
        mentions = _date_mentions(normalized, rolling_year)
        full_date_years = [
            int(match.group("year"))
            for match in re.finditer(
                r"(?P<year>20\d{2})[年/\-.]\d{1,2}[月/\-.]\d{1,2}日?",
                normalized,
            )
        ]
        if full_date_years:
            rolling_year = full_date_years[-1]
        for indicator, patterns in LAB_PATTERNS.items():
            for pattern in patterns:
                for match in re.finditer(pattern, normalized, flags=re.IGNORECASE):
                    value = float(match.group(1).replace(" ", ""))
                    preceding_dates = [item for position, item in mentions if position <= match.start()]
                    anchor_date = preceding_dates[-1] if preceding_dates else None
                    low, high = SAFETY_BOUNDS[indicator]
                    if low <= value <= high and (indicator, value) not in seen:
                        anchors.append(LabAnchor(indicator, value, paragraph[:300], anchor_date, len(anchors)))
                        seen.add((indicator, value))
    return anchors


def _extract_diagnosis(paragraphs: list[str]) -> str:
    for index, paragraph in enumerate(paragraphs):
        if paragraph in {"四、临床诊断", "临床诊断", "诊断", "（一）最终确诊诊断", "（一）最终诊断"}:
            diagnosis_lines: list[str] = []
            for candidate in paragraphs[index + 1:]:
                if re.match(r"^(?:五、|六、|七、|八、|治疗方案|疗效评估|随访转归|图片)", candidate):
                    break
                diagnosis_lines.append(candidate)
            return "；".join(diagnosis_lines)
    diagnosis_lines = [p for p in paragraphs if "诊断" in p][:4]
    return "；".join(diagnosis_lines)


def _has_affirmed_term(text: str, term: str) -> bool:
    for match in re.finditer(re.escape(term), text, flags=re.IGNORECASE):
        before = text[max(0, match.start() - 70):match.start()]
        after = text[match.end():match.end() + 28]
        if term.startswith("酒精") and before.endswith("非"):
            continue
        sentence_before = re.split(r"[；;。！？!?\n]", before)[-1]
        negations = list(re.finditer(r"否认|排除|无|未见|不支持|不考虑", sentence_before))
        affirmations = list(re.finditer(r"既往有|诊断为|确诊为|明确为|提示|证实|阳性|病史\s*\d*\s*年", sentence_before))
        last_negation = negations[-1].start() if negations else -1
        last_affirmation = affirmations[-1].start() if affirmations else -1
        if last_negation > last_affirmation and len(sentence_before) - last_negation <= 55:
            continue
        if re.match(r"^[\s，,：:（(]*(?:阴性|正常|已排除|可排除|不支持|不考虑|不除外|待排|[？?])", after):
            continue
        if (
            re.match(r"^[^；;。！？!?\n]{0,35}(?:阴性|排除|可能|待排|不除外)", after)
            and not re.match(r"^[^；;。！？!?\n]{0,35}(?:病史\s*\d+\s*年|[^，,；;。]{0,20}阳性)", after)
        ):
            continue
        return True
    return False


def _confirmed_diagnosis_scope(diagnosis_text: str, full_text: str) -> str:
    source = diagnosis_text or full_text[:2500]
    return re.split(
        r"鉴别诊断|五、治疗|治疗方案|六、疗效|七、图片",
        source,
        maxsplit=1,
    )[0].lower()


def _is_suspected_diagnosis_mention(text: str, start: int, end: int) -> bool:
    before = re.split(r"[；;。！？!?\n]", text[max(0, start - 45):start])[-1]
    after = text[end:end + 30]
    if re.match(r"^[\s（(：:，,、]*(?:不除外|待排(?:除)?|待除外|可疑|可能|[？?])", after):
        return True
    return bool(re.search(r"(?:可能|不除外|可疑|或)[^；;。！？!?\n]{0,18}(?:进展为|发展为)[^；;。！？!?\n]{0,8}$", before))


def _has_non_suspected_term(text: str, terms: Iterable[str]) -> bool:
    for term in terms:
        for match in re.finditer(re.escape(term), text, flags=re.IGNORECASE):
            if _has_affirmed_term(text, term) and not _is_suspected_diagnosis_mention(
                text, match.start(), match.end()
            ):
                return True
    return False


def _explicit_hcc_evidence(case: CaseRecord) -> bool:
    diagnostic_scope = _confirmed_diagnosis_scope(case.diagnosis_text, case.full_text)
    hcc_terms = ("肝细胞性肝癌", "肝细胞癌", "肝细胞肝癌", "肝癌", "hcc")
    if _has_non_suspected_term(diagnostic_scope, hcc_terms):
        return True
    pathological_patterns = [
        r"病理[^。；\n]{0,100}(?:肝细胞性?肝癌|肝细胞癌)",
        r"病理[^。；\n]{0,100}\bhcc\b",
        r"(?:肝细胞性?肝癌|肝细胞癌)[^。；\n]{0,80}(?:病理|免疫组化|分期)",
        r"(?:确诊|诊断为|证实)[^。；\n]{0,40}(?:肝细胞性?肝癌|肝细胞癌|\bhcc\b)",
        r"\bhcc\b[^。；\n]{0,80}(?:确诊|病灶|肿瘤|门静脉侵犯|卫星结节)",
    ]
    for pattern in pathological_patterns:
        for match in re.finditer(pattern, case.full_text, flags=re.IGNORECASE):
            for term in hcc_terms:
                for term_match in re.finditer(re.escape(term), match.group(0), flags=re.IGNORECASE):
                    start = match.start() + term_match.start()
                    end = match.start() + term_match.end()
                    if not _is_suspected_diagnosis_mention(case.full_text, start, end):
                        return True
    return False


def _explicit_cirrhosis_evidence(case: CaseRecord) -> bool:
    if _explicit_hcc_evidence(case):
        return False
    diagnostic_scope = _confirmed_diagnosis_scope(case.diagnosis_text, case.full_text)
    if _has_non_suspected_term(diagnostic_scope, ("肝硬化",)):
        return True
    for term in re.finditer("肝硬化", case.full_text):
        if _is_suspected_diagnosis_mention(case.full_text, term.start(), term.end()):
            continue
        before = re.split(r"[；;。！？!?\n]", case.full_text[max(0, term.start() - 90):term.start()])[-1]
        after = case.full_text[term.end():term.end() + 70]
        if re.search(r"(?:明确诊断为|最终诊断为|进展为|发展为)[^；;。！？!?\n]{0,60}$", before):
            return True
        if re.match(r"[^；;。！？!?\n]{0,60}(?:确诊|诊断明确)", after):
            return True
    return False


def _clinical_history_scope(full_text: str) -> str:
    sections: list[str] = []
    for pattern in (
        r"一、临床问诊(.*?)(?=二、全套实验室检查|二、实验室检查|二、辅助检查)",
        r"(?:个人生活史|个人史)[：:](.*?)(?=家族史|体格检查|二、|三、|四、)",
        r"既往史[：:](.*?)(?=个人生活史|个人史|家族史|体格检查|二、|三、|四、)",
    ):
        sections.extend(re.findall(pattern, full_text, flags=re.DOTALL))
    return "；".join(sections).lower()


def _has_affirmed_alcohol_exposure(text: str) -> bool:
    direct_terms = (
        "酒精性肝病", "酒精性脂肪肝", "酒精依赖", "酒精引起的戒断",
        "长期大量饮酒", "大量饮酒史",
    )
    if any(_has_affirmed_term(text, term) for term in direct_terms):
        return True
    patterns = [
        r"饮酒史\s*\d+\s*余?年[^。；\n]{0,80}(?:多饮白酒|(?:每日|每天|每周)[^。；\n]{0,35}(?:饮酒|白酒|啤酒)|\d+(?:\.\d+)?\s*(?:~|～|至|-)?\s*\d*\s*(?:g|ml|两)\s*/\s*(?:次|d|日|周))",
        r"多饮白酒",
        r"(?:每日|每天)(?:饮酒|饮啤酒|饮白酒)",
        r"饮酒量\s*\d+(?:\.\d+)?\s*(?:~|～|至|-)?\s*\d*\s*g\s*/\s*(?:次|d|日)",
        r"酒精及代谢双因素所致",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            sentence_before = re.split(r"[；;。！？!?\n]", text[max(0, match.start() - 55):match.start()])[-1]
            if re.search(r"(?:否认|无)[^；;。！？!?\n]{0,30}$", sentence_before):
                continue
            return True
    return False


def _classify_cohort(
    full_text: str, diagnosis_text: str, source_case_id: str = ""
) -> tuple[str, float, list[str]]:
    primary = (diagnosis_text or full_text[:2500]).lower()
    whole = full_text.lower()
    primary_main = re.split(r"鉴别诊断|需排除|可排除", primary, maxsplit=1)[0]
    diagnostic_scope = _confirmed_diagnosis_scope(diagnosis_text, full_text)
    history_scope = _clinical_history_scope(full_text)
    competing_scope = f"{diagnostic_scope}；{history_scope}"
    fatty_terms = [
        "脂肪肝", "脂肪性肝病", "脂肪性肝炎", "非酒精性脂肪", "代谢相关性脂肪",
        "代谢相关脂肪", "nafld", "nash", "mafld", "masld", "mash",
    ]
    competing_terms = {
        "competing_alcohol_related_liver_disease": [],
        "competing_viral_hepatitis": [
            "慢性乙型肝炎", "乙型病毒性肝炎", "慢性丙型肝炎", "病毒性肝炎", "丙肝后", "hbv-dna",
        ],
        "competing_drug_induced_liver_injury": [
            "药物性肝损伤", "药物诱导的自身免疫性肝炎", "药物诱导脂肪肝", "他莫昔芬致",
        ],
        "competing_autoimmune_hepatitis": ["自身免疫性肝炎"],
        "competing_liver_abscess": ["肝脓肿"],
        "competing_pregnancy_acute_fatty_liver": ["妊娠期急性脂肪肝", "aflp"],
        "competing_genetic_or_mitochondrial_disease": [
            "肝豆状核变性", "遗传代谢性疾病", "线粒体dna耗竭", "mngie", "madd", "hmgcs2",
            "prader-willi", "原发性肉碱缺乏症",
        ],
        "competing_non_hcc_liver_cancer": ["胆管细胞癌"],
    }
    metabolic_terms = {
        "metabolic_comorbidity_diabetes": ["2型糖尿病", "t2dm", "糖尿病前期", "胰岛素抵抗"],
        "metabolic_comorbidity_obesity": ["肥胖", "中心性肥胖", "腹型肥胖"],
        "metabolic_comorbidity_hypertension": ["高血压"],
        "metabolic_comorbidity_dyslipidemia": ["高脂血症", "血脂异常", "高甘油三酯"],
        "metabolic_comorbidity_metabolic_syndrome": ["代谢综合征"],
        "metabolic_comorbidity_pcos": ["多囊卵巢综合征"],
        "metabolic_comorbidity_sleep_apnea": ["睡眠呼吸暂停"],
    }
    reasons: list[str] = []
    has_fatty_evidence = any(_has_affirmed_term(whole, term) for term in fatty_terms)
    if has_fatty_evidence:
        reasons.append("documented_fatty_liver")
    else:
        reasons.append("insufficient_fatty_liver_evidence")
    for reason, terms in metabolic_terms.items():
        if any(_has_affirmed_term(whole, term) for term in terms):
            reasons.append(reason)
    competing_reasons = []
    if _has_affirmed_alcohol_exposure(competing_scope):
        competing_reasons.append("competing_alcohol_related_liver_disease")
    competing_reasons.extend(
        reason for reason, terms in competing_terms.items()
        if reason != "competing_alcohol_related_liver_disease"
        and any(_has_affirmed_term(competing_scope, term) for term in terms)
    )
    reasons.extend(competing_reasons)
    fatty_positions = [whole.find(term) for term in fatty_terms if whole.find(term) >= 0]
    advanced_positions = [whole.find(term) for term in ("肝硬化", "肝癌", "肝细胞癌") if whole.find(term) >= 0]
    documented_prephase = bool(fatty_positions and advanced_positions and min(fatty_positions) < min(advanced_positions)) \
        or "masld/mash 阶段" in whole \
        or bool(re.search(r"(?:脂肪肝|脂肪性肝病|nafld|mafld|masld|nash|mash).{0,50}(?:进展|发展|演变).{0,80}(?:肝硬化|肝癌|肝细胞癌)", whole))
    no_fatty_prephase = (
        any(_has_affirmed_term(primary_main, term) for term in ("肝硬化", "肝癌", "肝细胞癌"))
        and not documented_prephase
        and not re.search(r"(?:脂肪肝|脂肪性肝病).{0,30}(?:年|月|前).{0,80}(?:肝硬化|肝癌|肝细胞癌)", whole)
        and not re.search(r"(?:肝硬化|肝癌|肝细胞癌).{0,80}(?:既往|此前|多年|早期).{0,40}(?:脂肪肝|脂肪性肝病)", whole)
    )
    if no_fatty_prephase:
        reasons.append("advanced_stage_without_fatty_prephase")
    hcc_first_presentation_without_progression = (
        bool(re.search(r"(?:病理[^。；\n]{0,100})?(?:肝细胞性?肝癌|肝细胞癌)", whole))
        and "masld/mash 阶段" not in whole
        and not re.search(
            r"(?:脂肪肝|脂肪性肝病|nafld|mafld|masld|nash|mash).{0,80}(?:进展|发展|演变).{0,100}(?:肝癌|肝细胞癌)",
            whole,
        )
    )
    if hcc_first_presentation_without_progression:
        reasons.append("hcc_first_presentation_without_documented_progression")
    if source_case_id == "A6-1" or "nafld/nash的两个治疗策略" in whole:
        reasons.append("insufficient_patient_level_narrative")
    excluded = bool(
        competing_reasons
        or no_fatty_prephase
        or hcc_first_presentation_without_progression
        or "insufficient_fatty_liver_evidence" in reasons
        or "insufficient_patient_level_narrative" in reasons
    )
    positive = sum(term in whole for term in fatty_terms)
    score = positive * 2.5 - len(competing_reasons) * 4.0 - (4.0 if no_fatty_prephase else 0.0)
    if not excluded:
        reasons.append("eligible_no_competing_etiology")
        return "fatty_liver_progression", score, reasons
    return "mixed", score, reasons


def parse_case_documents(doc_a: Path, doc_b: Path) -> list[CaseRecord]:
    raw_cases = _segment_document("A", Path(doc_a)) + _segment_document("B", Path(doc_b))
    retained: list[CaseRecord] = []
    for raw in raw_cases:
        source_case_id = f"{raw['source']}{raw['source_number']}-{raw['source_occurrence']}"
        if source_case_id in EXCLUDED_SOURCE_CASES:
            continue
        paragraphs = raw["paragraphs"]
        full_text = "\n".join(paragraphs)
        diagnosis = _extract_diagnosis(paragraphs)
        cohort, score, classification_reasons = _classify_cohort(full_text, diagnosis, source_case_id)
        retained.append(CaseRecord(
            patient_id=f"P{len(retained) + 1:03d}",
            source=raw["source"],
            source_number=raw["source_number"],
            source_occurrence=raw["source_occurrence"],
            source_case_id=source_case_id,
            paragraphs=paragraphs,
            full_text=full_text,
            age=_extract_age(paragraphs),
            sex=_extract_sex(paragraphs),
            explicit_dates=_extract_dates(full_text),
            lab_anchors=_extract_lab_anchors(paragraphs),
            diagnosis_text=diagnosis,
            cohort_group=cohort,
            classification_reasons=classification_reasons,
            evidence_score=score,
        ))
    if len(retained) != 150:
        raise ValueError(f"Expected 150 retained cases, found {len(retained)}")
    return retained


def _add_months(value: date, months: int) -> date:
    total = value.year * 12 + value.month - 1 + months
    year, month_index = divmod(total, 12)
    month = month_index + 1
    day = min(value.day, [31, 29 if year % 4 == 0 and (year % 100 != 0 or year % 400 == 0) else 28,
                          31, 30, 31, 30, 31, 31, 30, 31, 30, 31][month - 1])
    return date(year, month, day)


def _deterministic_demographics(case: CaseRecord) -> tuple[int, str]:
    digest = int(hashlib.sha256(case.source_case_id.encode("utf-8")).hexdigest()[:12], 16)
    age = case.age if case.age is not None else 16 + digest % 70
    age = min(85, max(1, age))
    sex = case.sex if case.sex is not None else ("male" if digest % 2 == 0 else "female")
    return age, sex


def _risk_score(case: CaseRecord, age: int, sex: str) -> float:
    text = case.full_text.lower()
    score = case.evidence_score + (age - 40) * 0.05 + (0.5 if sex == "male" else 0.0)
    for term, weight in {
        "肝癌": 18, "肝细胞癌": 20, "hcc": 18, "肝硬化": 12, "纤维化高风险": 7,
        "血小板减少": 5, "腹水": 4, "门静脉": 3, "糖尿病": 2, "重度脂肪肝": 2,
    }.items():
        if term in text:
            score += weight
    if case.cohort_group == "fatty_liver_progression":
        score += 5
    return score


def _allocate_stages(cases: list[CaseRecord], config: GenerationConfig = GenerationConfig()) -> dict[str, str]:
    demographics = {case.patient_id: _deterministic_demographics(case) for case in cases}
    ranked = sorted(
        cases,
        key=lambda case: (-_risk_score(case, *demographics[case.patient_id]), case.patient_id),
    )
    hcc_ranked = sorted(
        ranked,
        key=lambda case: (
            -(20 if any(term in case.full_text.lower() for term in ("肝细胞癌", "肝癌", "hcc")) else 0)
            - _risk_score(case, *demographics[case.patient_id]),
            case.patient_id,
        ),
    )
    progression_cases = [case for case in cases if case.cohort_group == "fatty_liver_progression"]
    progression_ranked = sorted(
        progression_cases,
        key=lambda case: (-_risk_score(case, *demographics[case.patient_id]), case.patient_id),
    )
    target_event_count = min(max(20, len(progression_cases) - 6), len(progression_cases) - 1)
    forced_target_events = {case.patient_id for case in progression_ranked[:target_event_count]}
    protected_controls = {case.patient_id for case in progression_ranked[target_event_count:]}
    explicit_hcc = {case.patient_id for case in cases if _explicit_hcc_evidence(case)}
    explicit_cirrhosis = {case.patient_id for case in cases if _explicit_cirrhosis_evidence(case)}
    hcc_candidates = (
        [case.patient_id for case in hcc_ranked if case.patient_id in explicit_hcc]
        + [
            case.patient_id for case in hcc_ranked
            if case.patient_id in forced_target_events
            and case.patient_id not in explicit_hcc | explicit_cirrhosis
        ]
        + [
            case.patient_id for case in hcc_ranked
            if case.patient_id not in protected_controls | forced_target_events | explicit_hcc | explicit_cirrhosis
        ]
        + [
            case.patient_id for case in hcc_ranked
            if case.patient_id in protected_controls
            and case.patient_id not in explicit_hcc | explicit_cirrhosis
        ]
    )
    hcc_ids = set(hcc_candidates[:max(config.hcc_count, len(explicit_hcc))])
    remaining = [case for case in ranked if case.patient_id not in hcc_ids]
    cirrhosis_candidates = (
        [case.patient_id for case in remaining if case.patient_id in explicit_cirrhosis]
        + [
            case.patient_id for case in remaining
            if case.patient_id in forced_target_events and case.patient_id not in explicit_cirrhosis
        ]
        + [
            case.patient_id for case in remaining
            if case.patient_id not in protected_controls | forced_target_events | explicit_cirrhosis
        ]
        + [
            case.patient_id for case in remaining
            if case.patient_id in protected_controls and case.patient_id not in explicit_cirrhosis
        ]
    )
    cirrhosis_ids = set(cirrhosis_candidates[:max(config.cirrhosis_count, len(explicit_cirrhosis))])
    return {
        case.patient_id: ("hcc" if case.patient_id in hcc_ids else
                          "cirrhosis" if case.patient_id in cirrhosis_ids else "fatty_liver")
        for case in cases
    }


def _allocate_paths(cases: list[CaseRecord], stages: dict[str, str]) -> dict[str, str]:
    progressors = [case for case in cases if stages[case.patient_id] != "fatty_liver"]
    progressors.sort(key=lambda case: (-_risk_score(case, *_deterministic_demographics(case)), case.patient_id))
    paths: dict[str, str] = {}
    def anchor_count(case: CaseRecord, indicator: str) -> int:
        return len({anchor.value for anchor in case.lab_anchors if anchor.indicator == indicator})

    r1_eligible = [
        case for case in progressors
        if _deterministic_demographics(case)[0] > 50
        and _deterministic_demographics(case)[1] == "male"
        and anchor_count(case, "hba1c") <= 1
        and not any(anchor.anchor_date for anchor in case.lab_anchors if anchor.indicator == "hba1c")
        and anchor_count(case, "plt") == 0
    ]
    combined_count = min(len(r1_eligible), 12, max(5, (len(r1_eligible) + 1) // 2))
    combined = r1_eligible[:combined_count]
    r1_only = r1_eligible[combined_count:]
    assigned = {case.patient_id for case in combined + r1_only}
    remaining = [case for case in progressors if case.patient_id not in assigned]
    remaining.sort(key=lambda case: (anchor_count(case, "afp") > 0,
                                     -_risk_score(case, *_deterministic_demographics(case)), case.patient_id))
    r2_only = remaining[:15]
    for case in combined:
        paths[case.patient_id] = "r1_r2"
    for case in r1_only:
        paths[case.patient_id] = "r1"
    for case in r2_only:
        paths[case.patient_id] = "r2"
    for case in progressors:
        paths.setdefault(case.patient_id, "non_rule_progression")
    for case in cases:
        paths.setdefault(case.patient_id, "stable")
    return paths


def _select_explicit_dates(values: list[date], capacity: int) -> list[date]:
    unique = sorted(set(values))
    if not unique:
        return []
    best: list[date] = []
    for start_index in range(len(unique)):
        window = [value for value in unique[start_index:] if (value - unique[start_index]).days <= 1830]
        if len(window) > len(best) or (len(window) == len(best) and window[-1] > best[-1] if best else True):
            best = window
    if len(best) <= capacity:
        return best
    indices = np.linspace(0, len(best) - 1, capacity).round().astype(int)
    return [best[index] for index in sorted(set(indices))]


def _prioritize_case_dates(case: CaseRecord, values: list[date], capacity: int) -> list[date]:
    if capacity <= 0:
        return []
    anchor_indicators: dict[date, set[str]] = defaultdict(set)
    for anchor in case.lab_anchors:
        if anchor.anchor_date is not None:
            anchor_indicators[anchor.anchor_date].add(anchor.indicator)
    ranked = sorted(
        set(values),
        key=lambda value: (len(anchor_indicators[value]), value),
        reverse=True,
    )
    return sorted(ranked[:capacity])


def _choose_timeline(case: CaseRecord, rng: np.random.Generator, visit_count: int) -> list[date]:
    source_dates = [value for value in case.explicit_dates if value <= DATA_CUTOFF]
    source_dates.extend(
        anchor.anchor_date
        for anchor in case.lab_anchors
        if anchor.anchor_date is not None and anchor.anchor_date <= DATA_CUTOFF
    )
    explicit = _select_explicit_dates(source_dates, len(set(source_dates)))
    if explicit:
        earliest, latest = explicit[0], explicit[-1]
        if (latest - earliest).days >= 730:
            start, last = earliest, latest
        else:
            before = _add_months(latest, -24)
            after = _add_months(earliest, 24)
            if before >= date(2010, 1, 1):
                start, last = before, latest
            else:
                start, last = earliest, min(after, DATA_CUTOFF)
            if (last - start).days < 730:
                start, last = _add_months(DATA_CUTOFF, -24), DATA_CUTOFF
    else:
        year = int(rng.integers(2014, 2023))
        start = date(year, int(rng.integers(1, 13)), int(rng.integers(5, 24)))
        max_span = min(60, max(24, (DATA_CUTOFF.year - start.year) * 12 + DATA_CUTOFF.month - start.month))
        span_months = int(rng.integers(24, max_span + 1)) if max_span > 24 else 24
        last = min(_add_months(start, span_months), DATA_CUTOFF)
    if (last - start).days > 1830:
        start = _add_months(last, -60)
    explicit = [value for value in explicit if start <= value <= last and value not in {start, last}]
    explicit = _prioritize_case_dates(case, explicit, visit_count - len({start, last}))
    required = {start, last, *explicit}
    span_days = (last - start).days
    candidates = [start + timedelta(days=int(round(span_days * fraction)))
                  for fraction in np.linspace(0.0, 1.0, visit_count * 3)]
    dates = sorted(required)
    for candidate in candidates:
        if len(dates) >= visit_count:
            break
        if candidate not in dates:
            dates.append(candidate)
            dates.sort()
    while len(dates) < visit_count:
        gaps = [(dates[index + 1] - dates[index]).days for index in range(len(dates) - 1)]
        gap_index = int(np.argmax(gaps))
        candidate = dates[gap_index] + timedelta(days=max(1, gaps[gap_index] // 2))
        if candidate not in dates:
            dates.append(candidate)
            dates.sort()
    return dates[:visit_count]


def _anchor_baselines(case: CaseRecord) -> dict[str, float]:
    by_indicator: dict[str, list[float]] = defaultdict(list)
    for anchor in case.lab_anchors:
        by_indicator[anchor.indicator].append(anchor.value)
    result = dict(DEFAULT_BASELINES)
    for indicator, values in by_indicator.items():
        result[indicator] = float(values[0])
    text = case.full_text.lower()
    if "肥胖" in text or "中心性肥胖" in text:
        result["bmi"] = max(result["bmi"], 29.0)
        result["waist"] = max(result["waist"], 98.0)
    if "糖尿病" in text:
        result["hba1c"] = max(result["hba1c"], 7.2)
    if "血小板增多" in text:
        result["plt"] = max(result["plt"], 500.0)
    return result


def _round_value(indicator: str, value: float) -> int | float:
    low, high = SAFETY_BOUNDS[indicator]
    value = min(high, max(low, value))
    if indicator in {"alt", "ast", "ggt", "plt", "waist"}:
        return round(value, 1)
    return round(value, 2)


def _anchor_assignments(case: CaseRecord, timeline: list[date]) -> tuple[dict[tuple[int, str], float], list[dict[str, Any]]]:
    by_indicator: dict[str, list[LabAnchor]] = defaultdict(list)
    for anchor in case.lab_anchors:
        if not any(math.isclose(anchor.value, existing.value) for existing in by_indicator[anchor.indicator]):
            by_indicator[anchor.indicator].append(anchor)
    assignments: dict[tuple[int, str], float] = {}
    conflicts: list[dict[str, Any]] = []
    for indicator, anchors in by_indicator.items():
        if len(anchors) > len(timeline):
            conflicts.append({
                "patient_id": case.patient_id,
                "indicator": indicator,
                "reason": "more_unique_source_values_than_visits",
                "source_values": [anchor.value for anchor in anchors],
            })
        dated = [anchor for anchor in anchors if anchor.anchor_date in timeline]
        omitted_dated = [
            anchor for anchor in anchors
            if anchor.anchor_date is not None and anchor.anchor_date not in timeline
        ]
        for anchor in omitted_dated:
            conflicts.append({
                "patient_id": case.patient_id,
                "indicator": indicator,
                "reason": "dated_source_anchor_not_on_timeline",
                "source_value": anchor.value,
                "source_date": anchor.anchor_date.isoformat(),
            })
        undated = [anchor for anchor in anchors if anchor.anchor_date is None]
        undated.extend(omitted_dated)
        used_indices: set[int] = set()
        for anchor in dated:
            index = timeline.index(anchor.anchor_date)  # exact source date is a visit node
            if (index, indicator) in assignments:
                conflicts.append({
                    "patient_id": case.patient_id, "indicator": indicator,
                    "reason": "multiple_source_values_same_date", "source_value": anchor.value,
                    "source_date": anchor.anchor_date.isoformat(),
                })
                undated.append(anchor)
            else:
                assignments[(index, indicator)] = anchor.value
                used_indices.add(index)
        available = [index for index in range(len(timeline)) if index not in used_indices]
        for anchor, index in zip(undated, available):
            assignments[(index, indicator)] = anchor.value
            used_indices.add(index)
        for anchor in undated[len(available):]:
            conflicts.append({
                "patient_id": case.patient_id, "indicator": indicator,
                "reason": "source_value_not_representable", "source_value": anchor.value,
                "source_context": anchor.context,
            })
    return assignments, conflicts


def _generate_values(
    case: CaseRecord,
    path: str,
    stage: str,
    visit_count: int,
    timeline: list[date],
    rng: np.random.Generator,
) -> tuple[list[dict[str, int | float | str]], list[dict[str, Any]]]:
    base = _anchor_baselines(case)
    patient_scale = {indicator: float(rng.normal(1.0, 0.045)) for indicator in INDICATORS}
    rows: list[dict[str, int | float | str]] = []
    for index in range(visit_count):
        progress = index / max(1, visit_count - 1)
        late = max(0.0, (progress - 0.45) / 0.55)
        values: dict[str, float] = {}
        inflammation = 1.0 + (0.65 * progress if stage != "fatty_liver" else -0.20 * progress)
        values["alt"] = base["alt"] * inflammation + rng.normal(0, 7 + 12 * progress)
        values["ast"] = base["ast"] * (1.0 + (0.85 * progress if stage != "fatty_liver" else -0.12 * progress)) + rng.normal(0, 6)
        values["ggt"] = base["ggt"] * (1.0 + (0.75 * progress if stage != "fatty_liver" else -0.10 * progress)) + rng.normal(0, 9)
        values["tbil"] = base["tbil"] + (16 * late if stage != "fatty_liver" else rng.normal(0, 1.2)) + rng.normal(0, 1.5)
        values["alb"] = base["alb"] - (7.5 * progress if stage != "fatty_liver" else 0.8 * progress) + rng.normal(0, 0.9)
        values["plt"] = base["plt"] + rng.normal(0, max(7, base["plt"] * 0.035))
        values["hba1c"] = base["hba1c"] + rng.normal(0, 0.12)
        values["afp"] = base["afp"] + rng.normal(0, max(0.25, base["afp"] * 0.05))
        values["waist"] = base["waist"] + (-4.0 * progress if stage == "fatty_liver" else 2.5 * progress) + rng.normal(0, 1.0)
        values["bmi"] = base["bmi"] + (-1.5 * progress if stage == "fatty_liver" else 0.7 * progress) + rng.normal(0, 0.35)

        if path in {"r1", "r1_r2"}:
            values["hba1c"] = base["hba1c"] + 0.48 * index + rng.normal(0, 0.06)
            values["plt"] = base["plt"] * (1.0 - 0.10 * index) + rng.normal(0, max(4.0, base["plt"] * 0.018))
        elif path == "non_rule_progression":
            values["hba1c"] = base["hba1c"] + 0.10 * index + rng.normal(0, 0.18)
            values["plt"] = base["plt"] * (1.0 - 0.035 * index) + rng.normal(0, 10)

        if path in {"r2", "r1_r2"}:
            afp_start = max(0, visit_count - 3)
            values["afp"] = base["afp"] + max(0, index - afp_start + 1) * (4.5 if stage == "hcc" else 2.3) + rng.normal(0, 0.18)
        elif stage == "hcc":
            values["afp"] = base["afp"] + 1.5 * late + rng.normal(0, 0.6)

        row: dict[str, int | float | str] = {}
        for indicator, value in values.items():
            row[indicator] = _round_value(indicator, value * patient_scale[indicator])
        rows.append(row)

    non_core_missing = {"alt": 0.08, "ast": 0.08, "ggt": 0.12, "tbil": 0.10, "alb": 0.10, "waist": 0.24, "bmi": 0.16}
    anchored_indicators = {anchor.indicator for anchor in case.lab_anchors}
    for row_index, row in enumerate(rows):
        for indicator, probability in non_core_missing.items():
            if indicator in anchored_indicators and row_index == 0:
                continue
            if rng.random() < probability:
                row[indicator] = ""
    assignments, conflicts = _anchor_assignments(case, timeline)
    for (row_index, indicator), value in assignments.items():
        rows[row_index][indicator] = round(float(value), 2)
    assigned_indicators = {indicator for _, indicator in assignments}
    if path in {"r1", "r1_r2"}:
        fixed_hba_indices = {index for index, indicator in assignments if indicator == "hba1c"}
        for index in range(1, len(rows)):
            if index not in fixed_hba_indices:
                rows[index]["hba1c"] = round(
                    max(float(rows[index]["hba1c"]), float(rows[index - 1]["hba1c"]) + 0.25), 2
                )
        if "plt" not in assigned_indicators:
            plt_baseline = float(np.mean([float(rows[0]["plt"]), float(rows[1]["plt"])]))
            rows[-1]["plt"] = round(min(float(rows[-1]["plt"]), plt_baseline * 0.75), 2)
    if path in {"r2", "r1_r2"} and "afp" not in assigned_indicators:
        for index in range(max(1, len(rows) - 2), len(rows)):
            rows[index]["afp"] = round(max(float(rows[index]["afp"]), float(rows[index - 1]["afp"]) + 1.0), 2)
    return rows, conflicts


def _date_index_for_stage(stage: str, visit_count: int) -> tuple[int | None, int | None]:
    if stage == "fatty_liver":
        return None, None
    if stage == "cirrhosis":
        return max(1, visit_count - 2), None
    hcc_index = visit_count - 1
    cirrhosis_index = max(1, hcc_index - 1) if visit_count >= 4 else None
    return cirrhosis_index, hcc_index


def _fatty_liver_date(case: CaseRecord, timeline: list[date]) -> tuple[date, str]:
    fatty_terms = ("脂肪肝", "脂肪性肝病", "脂肪性肝炎", "nafld", "nash", "mafld", "masld", "mash")
    candidates: list[date] = []
    for paragraph in case.paragraphs:
        lowered = paragraph.lower()
        if any(_has_affirmed_term(lowered, term) for term in fatty_terms):
            candidates.extend(value for value in _extract_dates(paragraph) if value <= timeline[0])
    if candidates:
        return max(candidates), "source_anchored"
    earliest_encounter = min(case.explicit_dates) if case.explicit_dates else timeline[0]
    relative_patterns = [
        r"(\d{1,2})\s*(?:余|多)?\s*年前[^。；\n]{0,140}(?:发现|诊断|确诊|查出)[^。；\n]{0,60}(?:脂肪肝|脂肪性肝病|脂肪性肝炎|nafld|nash|mafld|masld|mash)",
        r"(?:脂肪肝|脂肪性肝病|脂肪性肝炎|nafld|nash|mafld|masld|mash)(?:病史)?\s*(\d{1,2})\s*(?:余|多)?\s*年",
        r"(?:发现|诊断|确诊|查出)[^。；\n]{0,35}(?:脂肪肝|脂肪性肝病|脂肪性肝炎|nafld|nash|mafld|masld|mash)\s*(\d{1,2})\s*(?:余|多)?\s*年",
    ]
    relative_years: list[int] = []
    lowered = case.full_text.lower()
    for pattern in relative_patterns:
        relative_years.extend(int(value) for value in re.findall(pattern, lowered, flags=re.IGNORECASE))
    if relative_years:
        inferred = _add_months(earliest_encounter, -12 * max(relative_years))
        if inferred < timeline[-1] and (timeline[-1] - inferred).days >= 730:
            return inferred, "source_relative_history"
    return timeline[0], "generated_baseline"


def generate_dataset(
    cases: list[CaseRecord], config: GenerationConfig = GenerationConfig()
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    if len(cases) != 150:
        raise ValueError("Generation requires exactly 150 retained cases")
    rng = np.random.default_rng(config.seed)
    stages = _allocate_stages(cases, config)
    paths = _allocate_paths(cases, stages)
    patients: list[dict[str, Any]] = []
    visits: list[dict[str, Any]] = []
    anchor_conflicts: list[dict[str, Any]] = []
    anchor_counts = Counter()
    fatty_liver_date_sources: dict[str, str] = {}
    for case in cases:
        age, sex = _deterministic_demographics(case)
        stage = stages[case.patient_id]
        path = paths[case.patient_id]
        max_unique_anchor_values = max(
            (len({anchor.value for anchor in case.lab_anchors if anchor.indicator == indicator})
             for indicator in INDICATORS),
            default=0,
        )
        minimum_visits = max(config.min_visits, min(config.max_visits, max_unique_anchor_values))
        visit_count = int(rng.integers(minimum_visits, config.max_visits + 1))
        timeline = _choose_timeline(case, rng, visit_count)
        values, conflicts = _generate_values(case, path, stage, visit_count, timeline, rng)
        anchor_conflicts.extend(conflicts)
        cirrhosis_index, hcc_index = _date_index_for_stage(stage, visit_count)
        cirrhosis_date = timeline[cirrhosis_index].isoformat() if cirrhosis_index is not None else ""
        hcc_date = timeline[hcc_index].isoformat() if hcc_index is not None else ""
        lost = "yes" if stage == "fatty_liver" and int(case.patient_id[1:]) % 13 == 0 else "no"
        fatty_liver_date, fatty_liver_date_source = _fatty_liver_date(case, timeline)
        fatty_liver_date_sources[case.patient_id] = fatty_liver_date_source
        patients.append({
            "patient_id": case.patient_id,
            "age": age,
            "sex": sex,
            "cohort_group": case.cohort_group,
            "fatty_liver_date": fatty_liver_date.isoformat(),
            "final_stage": stage,
            "cirrhosis_date": cirrhosis_date,
            "hcc_date": hcc_date,
            "last_followup_date": timeline[-1].isoformat(),
            "lost_to_followup": lost,
        })
        for visit_date, row_values in zip(timeline, values):
            visits.append({"patient_id": case.patient_id, "visit_date": visit_date.isoformat(), **row_values})
        anchor_counts.update(anchor.indicator for anchor in case.lab_anchors)
    report = _build_quality_report(
        cases, patients, visits, paths, anchor_counts, anchor_conflicts, fatty_liver_date_sources, config.seed
    )
    return patients, visits, report


def validate_dataset(patients: list[dict[str, Any]], visits: list[dict[str, Any]]) -> dict[str, Any]:
    errors: list[str] = []
    if len(patients) != 150:
        errors.append(f"patients row count is {len(patients)}, expected 150")
    expected_ids = [f"P{i:03d}" for i in range(1, 151)]
    patient_ids = [row.get("patient_id") for row in patients]
    if patient_ids != expected_ids:
        errors.append("patient_id values are not consecutive P001-P150")
    allowed = {
        "sex": {"male", "female"},
        "cohort_group": {"fatty_liver_progression", "mixed"},
        "final_stage": {"fatty_liver", "cirrhosis", "hcc"},
        "lost_to_followup": {"yes", "no"},
    }
    by_patient: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in visits:
        by_patient[row.get("patient_id", "")].append(row)
    if set(by_patient) - set(patient_ids):
        errors.append("visits contains orphan patient IDs")
    for patient in patients:
        pid = patient["patient_id"]
        if not 1 <= int(patient["age"]) <= 85:
            errors.append(f"{pid}: age out of range")
        for field_name, values in allowed.items():
            if patient[field_name] not in values:
                errors.append(f"{pid}: invalid {field_name}")
        patient_visits = by_patient.get(pid, [])
        if not 3 <= len(patient_visits) <= 6:
            errors.append(f"{pid}: visit count {len(patient_visits)} outside 3-6")
            continue
        dates = [date.fromisoformat(row["visit_date"]) for row in patient_visits]
        if dates != sorted(set(dates)):
            errors.append(f"{pid}: visit dates not strictly increasing")
        if (dates[-1] - dates[0]).days < 730:
            errors.append(f"{pid}: follow-up span below 24 months")
        if (dates[-1] - dates[0]).days > 1830:
            errors.append(f"{pid}: follow-up span above 60 months")
        if patient["last_followup_date"] != dates[-1].isoformat():
            errors.append(f"{pid}: last_followup_date mismatch")
        fatty = date.fromisoformat(patient["fatty_liver_date"])
        cirrhosis = date.fromisoformat(patient["cirrhosis_date"]) if patient["cirrhosis_date"] else None
        hcc = date.fromisoformat(patient["hcc_date"]) if patient["hcc_date"] else None
        stage = patient["final_stage"]
        if stage == "fatty_liver" and (cirrhosis or hcc):
            errors.append(f"{pid}: fatty_liver stage has event dates")
        if stage == "cirrhosis" and not (cirrhosis and fatty < cirrhosis <= dates[-1] and not hcc):
            errors.append(f"{pid}: invalid cirrhosis date logic")
        if stage == "hcc" and not (hcc and fatty < hcc <= dates[-1]):
            errors.append(f"{pid}: invalid hcc date logic")
        if cirrhosis and hcc and not cirrhosis < hcc:
            errors.append(f"{pid}: cirrhosis_date must precede hcc_date")
        for indicator in ("plt", "hba1c", "afp"):
            count = sum(row[indicator] != "" for row in patient_visits)
            if count < 3:
                errors.append(f"{pid}: {indicator} has only {count} values")
        for row in patient_visits:
            for indicator in INDICATORS:
                value = row[indicator]
                if value == "":
                    continue
                if not isinstance(value, (int, float)):
                    errors.append(f"{pid}: {indicator} is not numeric")
                    continue
                low, high = SAFETY_BOUNDS[indicator]
                if not low <= float(value) <= high:
                    errors.append(f"{pid}: {indicator}={value} outside bounds")
    stage_counts = Counter(row["final_stage"] for row in patients)
    if stage_counts != Counter({"fatty_liver": 75, "cirrhosis": 50, "hcc": 25}):
        errors.append(f"unexpected stage counts: {dict(stage_counts)}")
    cohort_stage = Counter((row["cohort_group"], row["final_stage"]) for row in patients)
    progression_events = cohort_stage[("fatty_liver_progression", "cirrhosis")] + cohort_stage[("fatty_liver_progression", "hcc")]
    progression_controls = cohort_stage[("fatty_liver_progression", "fatty_liver")]
    if progression_events < 20 or progression_controls == 0:
        errors.append("fatty_liver_progression must contain at least 20 events and non-progressing controls")
    return {"errors": errors, "error_count": len(errors)}


def _quantiles(values: Iterable[float]) -> dict[str, float] | None:
    array = np.asarray(list(values), dtype=float)
    if array.size == 0:
        return None
    return {key: round(float(value), 3) for key, value in zip(
        ("min", "p25", "median", "p75", "max"), np.quantile(array, [0, 0.25, 0.5, 0.75, 1])
    )}


def _actual_rule_signal_counts(
    patients: list[dict[str, Any]], visits: list[dict[str, Any]]
) -> dict[str, int]:
    patient_by_id = {row["patient_id"]: row for row in patients}
    by_patient: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in visits:
        by_patient[row["patient_id"]].append(row)
    r1_ids: set[str] = set()
    r2_ids: set[str] = set()
    for pid, rows in by_patient.items():
        patient = patient_by_id[pid]
        plt_values = [float(row["plt"]) for row in rows if row["plt"] != ""]
        hba_values = [float(row["hba1c"]) for row in rows if row["hba1c"] != ""]
        afp_values = [float(row["afp"]) for row in rows if row["afp"] != ""]
        hba_rises = len(hba_values) >= 3 and hba_values[-3] < hba_values[-2] < hba_values[-1]
        plt_drop = len(plt_values) >= 3 and plt_values[-1] <= 0.80 * np.mean(plt_values[:2])
        afp_rises = len(afp_values) >= 3 and afp_values[-3] < afp_values[-2] < afp_values[-1]
        if patient["sex"] == "male" and int(patient["age"]) > 50 and hba_rises and plt_drop:
            r1_ids.add(pid)
        if afp_rises:
            r2_ids.add(pid)
    return {
        "r1": len(r1_ids),
        "r2": len(r2_ids),
        "r1_r2": len(r1_ids & r2_ids),
        "neither": len(patient_by_id) - len(r1_ids | r2_ids),
    }


def _actual_rule_memberships(
    patients: list[dict[str, Any]], visits: list[dict[str, Any]]
) -> tuple[set[str], set[str]]:
    patient_by_id = {row["patient_id"]: row for row in patients}
    by_patient: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in visits:
        by_patient[row["patient_id"]].append(row)
    r1_ids: set[str] = set()
    r2_ids: set[str] = set()
    for pid, rows in by_patient.items():
        patient = patient_by_id[pid]
        plt_values = [float(row["plt"]) for row in rows if row["plt"] != ""]
        hba_values = [float(row["hba1c"]) for row in rows if row["hba1c"] != ""]
        afp_values = [float(row["afp"]) for row in rows if row["afp"] != ""]
        hba_rises = len(hba_values) >= 3 and hba_values[-3] < hba_values[-2] < hba_values[-1]
        plt_drop = len(plt_values) >= 3 and plt_values[-1] <= 0.80 * np.mean(plt_values[:2])
        afp_rises = len(afp_values) >= 3 and afp_values[-3] < afp_values[-2] < afp_values[-1]
        if patient["sex"] == "male" and int(patient["age"]) > 50 and hba_rises and plt_drop:
            r1_ids.add(pid)
        if afp_rises:
            r2_ids.add(pid)
    return r1_ids, r2_ids


def _assigned_path_mismatches(
    paths: dict[str, str], patients: list[dict[str, Any]], visits: list[dict[str, Any]]
) -> list[dict[str, str]]:
    r1_ids, r2_ids = _actual_rule_memberships(patients, visits)
    mismatches = []
    for pid, path in paths.items():
        if path in {"r1", "r1_r2"} and pid not in r1_ids:
            mismatches.append({"patient_id": pid, "assigned_path": path, "missing_rule": "r1"})
        if path in {"r2", "r1_r2"} and pid not in r2_ids:
            mismatches.append({"patient_id": pid, "assigned_path": path, "missing_rule": "r2"})
    return mismatches


def _cohort_rule_signal_counts(
    paths: dict[str, str], patients: list[dict[str, Any]], visits: list[dict[str, Any]]
) -> dict[str, dict[str, int]]:
    r1_ids, r2_ids = _actual_rule_memberships(patients, visits)
    patient_by_id = {row["patient_id"]: row for row in patients}
    result: dict[str, dict[str, int]] = {}
    for cohort in ("fatty_liver_progression", "mixed"):
        ids = {pid for pid, row in patient_by_id.items() if row["cohort_group"] == cohort}
        result[cohort] = {
            "r1": len(ids & r1_ids),
            "r2": len(ids & r2_ids),
            "r1_r2": len(ids & r1_ids & r2_ids),
            "neither": len(ids - (r1_ids | r2_ids)),
        }
    return result


def _build_quality_report(
    cases: list[CaseRecord],
    patients: list[dict[str, Any]],
    visits: list[dict[str, Any]],
    paths: dict[str, str],
    anchor_counts: Counter,
    anchor_conflicts: list[dict[str, Any]],
    fatty_liver_date_sources: dict[str, str],
    seed: int,
) -> dict[str, Any]:
    by_patient: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in visits:
        by_patient[row["patient_id"]].append(row)
    missing = {}
    quantiles = {}
    for indicator in INDICATORS:
        values = [float(row[indicator]) for row in visits if row[indicator] != ""]
        missing[indicator] = round(1.0 - len(values) / len(visits), 4)
        quantiles[indicator] = _quantiles(values)
    spans = [
        (date.fromisoformat(rows[-1]["visit_date"]) - date.fromisoformat(rows[0]["visit_date"])).days / 30.4375
        for rows in by_patient.values()
    ]
    visit_counts = [len(rows) for rows in by_patient.values()]
    patient_by_id = {row["patient_id"]: row for row in patients}
    outcome_assignment_audit: dict[str, dict[str, str]] = {}
    generated_outcome_ids: dict[str, list[str]] = {"fatty_liver": [], "cirrhosis": [], "hcc": []}
    for case in cases:
        stage = patient_by_id[case.patient_id]["final_stage"]
        if _explicit_hcc_evidence(case):
            source = "explicit_hcc"
        elif _explicit_cirrhosis_evidence(case):
            source = "explicit_cirrhosis"
        else:
            source = "generated_stage_assignment"
            generated_outcome_ids[stage].append(case.patient_id)
        outcome_assignment_audit[case.patient_id] = {
            "source_case_id": case.source_case_id,
            "final_stage": stage,
            "source": source,
        }
    return {
        "seed": seed,
        "generator_version": GENERATOR_VERSION,
        "patient_count": len(patients),
        "visit_count": len(visits),
        "source_counts": dict(Counter(case.source for case in cases)),
        "excluded_source_cases": sorted(EXCLUDED_SOURCE_CASES),
        "cohort_counts": dict(Counter(row["cohort_group"] for row in patients)),
        "cohort_classification_reasons": {
            case.patient_id: {
                "source_case_id": case.source_case_id,
                "cohort_group": case.cohort_group,
                "reasons": case.classification_reasons,
            }
            for case in cases
        },
        "classification_reason_counts": dict(Counter(
            reason for case in cases for reason in case.classification_reasons
        )),
        "stage_counts": dict(Counter(row["final_stage"] for row in patients)),
        "outcome_assignment_audit": outcome_assignment_audit,
        "generated_outcome_ids": generated_outcome_ids,
        "lost_to_followup_counts": dict(Counter(row["lost_to_followup"] for row in patients)),
        "generated_lost_to_followup_ids": [
            row["patient_id"] for row in patients if row["lost_to_followup"] == "yes"
        ],
        "source_anchored_fatty_liver_date_ids": sorted(
            patient_id for patient_id, source in fatty_liver_date_sources.items() if source == "source_anchored"
        ),
        "source_relative_fatty_liver_date_ids": sorted(
            patient_id for patient_id, source in fatty_liver_date_sources.items() if source == "source_relative_history"
        ),
        "generated_fatty_liver_date_ids": sorted(
            patient_id for patient_id, source in fatty_liver_date_sources.items() if source == "generated_baseline"
        ),
        "intended_use": [
            "data_import_pipeline_testing",
            "rule_mining_workflow_mechanism_validation",
            "ui_and_statistical_function_demonstration",
        ],
        "prohibited_uses": [
            "clinical_rule_discovery_claims",
            "real_world_clinical_evidence",
            "diagnosis_or_treatment_decisions",
            "unlabeled_clinical_research_source_data",
        ],
        "embedded_rule_paths": {
            "r1": "generated HbA1c-rise plus PLT-decline signal for workflow validation",
            "r2": "generated AFP-rise signal for workflow validation",
            "r1_r2": "generated combined R1 and R2 signal for workflow validation",
        },
        "path_counts": dict(Counter(paths.values())),
        "actual_rule_signal_counts": _actual_rule_signal_counts(patients, visits),
        "assigned_path_mismatches": _assigned_path_mismatches(paths, patients, visits),
        "cohort_rule_signal_counts": _cohort_rule_signal_counts(paths, patients, visits),
        "cohort_stage_cross_counts": {
            f"{cohort}|{stage}": count for (cohort, stage), count in sorted(
                Counter((row["cohort_group"], row["final_stage"]) for row in patients).items()
            )
        },
        "visit_count_distribution": dict(sorted(Counter(visit_counts).items())),
        "followup_months": _quantiles(spans),
        "missing_rates": missing,
        "indicator_quantiles": quantiles,
        "source_anchor_counts": dict(anchor_counts),
        "source_anchor_conflicts": anchor_conflicts,
        "validation": validate_dataset(patients, visits),
    }


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _csv_value(value: Any) -> Any:
    if value is None:
        return ""
    return value


def write_outputs(
    output_dir: Path,
    cases: list[CaseRecord],
    patients: list[dict[str, Any]],
    visits: list[dict[str, Any]],
    report: dict[str, Any],
    doc_a: Path,
    doc_b: Path,
) -> dict[str, Path]:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    validation = validate_dataset(patients, visits)
    if validation["errors"]:
        raise ValueError("Dataset validation failed: " + "; ".join(validation["errors"][:10]))
    patients_path = output_dir / "patients.csv"
    visits_path = output_dir / "visits.csv"
    quality_path = output_dir / "quality_report.json"
    provenance_path = output_dir / "DATA_PROVENANCE.md"
    extracted_cases_path = output_dir / "extracted_cases.json"
    with patients_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=PATIENT_HEADERS, lineterminator="\n")
        writer.writeheader()
        writer.writerows({key: _csv_value(row.get(key, "")) for key in PATIENT_HEADERS} for row in patients)
    with visits_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=VISIT_HEADERS, lineterminator="\n")
        writer.writeheader()
        writer.writerows({key: _csv_value(row.get(key, "")) for key in VISIT_HEADERS} for row in visits)
    final_report = dict(report)
    final_report["validation"] = validation
    quality_path.write_text(json.dumps(final_report, ensure_ascii=False, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    extracted_cases = []
    for case in cases:
        extracted_cases.append({
            "patient_id": case.patient_id,
            "source_case_id": case.source_case_id,
            "age": case.age,
            "sex": case.sex,
            "cohort_group": case.cohort_group,
            "classification_reasons": case.classification_reasons,
            "explicit_dates": [value.isoformat() for value in case.explicit_dates],
            "lab_anchors": [
                {
                    "indicator": anchor.indicator,
                    "value": anchor.value,
                    "anchor_date": anchor.anchor_date.isoformat() if anchor.anchor_date else None,
                    "sequence": anchor.sequence,
                    "context": anchor.context,
                }
                for anchor in case.lab_anchors
            ],
            "diagnosis_text": case.diagnosis_text,
        })
    extracted_cases_path.write_text(
        json.dumps(extracted_cases, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    mapping_lines = [
        f"| {case.patient_id} | {case.source_case_id} | {case.age if case.age is not None else '未明确'} | "
        f"{case.sex or '未明确'} | {case.cohort_group} | {len(case.lab_anchors)} |"
        for case in cases
    ]
    provenance = f"""# 数据来源与生成说明

## 生成信息

- 生成日期：2026-08-18
- 固定随机种子：`{report['seed']}`
- 生成器版本：`{GENERATOR_VERSION}`
- 生成器：`scripts/generate_fatty_liver_longitudinal.py`
- 患者数：{len(patients)}
- 随访记录数：{len(visits)}

## 输入文档

| 文档 | SHA-256 |
|---|---|
| `{Path(doc_a).name}` | `{_file_sha256(Path(doc_a))}` |
| `{Path(doc_b).name}` | `{_file_sha256(Path(doc_b))}` |

## 病例筛选

原始共解析 155 个病例段，固定排除 `A27-1`、`A30-1`、`A32-1`、`A34-1`、`A35-1`，保留 150 例。

## 字段来源

- 年龄、性别、明确日期、明确诊断及能够识别的检验值优先来自病例文本。
- 缺失人口学字段使用稳定规则补齐。
- 原病例没有提供的纵向日期与指标值由固定种子生成器补齐。
- `patients.csv` 与 `visits.csv` 保持采集规范的固定列，不增加来源或分类理由字段。
- `fatty_liver_progression` 允许糖尿病、肥胖、高血压、血脂异常等常见代谢共病；明确竞争性肝病病因、无脂肪肝证据或无脂肪肝前期的晚期首诊病例保留为 `mixed`。
- 每例分类理由记录在 `quality_report.json` 和 `extracted_cases.json`，用于逐例审计。
- 原文明确 HCC 或肝硬化时锁定对应结局；其余结局为达到方法验证所需事件量而生成分配。逐例来源见 `quality_report.json` 的 `outcome_assignment_audit` 和 `generated_outcome_ids`。
- `fatty_liver_date` 优先采用与脂肪肝证据相邻的明确日期；只有“X 年前发现/诊断脂肪肝”时，以病例中最早明确日期（若无则以生成首访）反推年份；仍无法定位时使用生成基线。三类患者清单见质量报告。
- `lost_to_followup=yes` 为固定规则生成的少量流程测试状态，不代表病例原文记载；患者清单见 `generated_lost_to_followup_ids`。
- 生成值不得表述为从已遗失的原始检验报告中补采得到，不得作为真实世界临床证据、诊疗依据或未经说明的临床研究原始数据。

## 使用边界

- 允许用途：数据导入管线测试、规律挖掘流程机制验证、界面与统计功能演示。
- 禁止用途：宣称发现真实临床规律、形成真实世界临床证据、支持诊疗决策或作为未经说明的临床研究原始数据。
- R1、R2、R1+R2 信号是生成器有意植入的流程验证路径；从本数据重新挖出这些信号属于机制回归测试，不是独立临床发现。

## 汇总

- 队列：{json.dumps(report['cohort_counts'], ensure_ascii=False)}
- 结局：{json.dumps(report['stage_counts'], ensure_ascii=False)}
- 路径：{json.dumps(report['path_counts'], ensure_ascii=False)}
- actual_rule_signal_counts：{json.dumps(report['actual_rule_signal_counts'], ensure_ascii=False)}
- 队列 × 结局：{json.dumps(report['cohort_stage_cross_counts'], ensure_ascii=False)}
- 各指标缺失率：{json.dumps(report['missing_rates'], ensure_ascii=False)}
- 原文锚点冲突数：{len(report['source_anchor_conflicts'])}
- 随访次数分布：{json.dumps(report['visit_count_distribution'], ensure_ascii=False)}

## 患者来源映射

| patient_id | 源病例段 | 原文年龄 | 原文性别 | 队列分组 | 提取检验锚点数 |
|---|---|---:|---|---|---:|
{chr(10).join(mapping_lines)}
"""
    provenance_path.write_text(provenance, encoding="utf-8")
    return {
        "patients": patients_path,
        "visits": visits_path,
        "quality": quality_path,
        "provenance": provenance_path,
        "extracted_cases": extracted_cases_path,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate case-constrained fatty-liver longitudinal data")
    parser.add_argument("--doc-a", type=Path, required=True)
    parser.add_argument("--doc-b", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    cases = parse_case_documents(args.doc_a, args.doc_b)
    patients, visits, report = generate_dataset(cases, GenerationConfig())
    paths = write_outputs(args.output_dir, cases, patients, visits, report, args.doc_a, args.doc_b)
    print(json.dumps({
        "patients": len(patients),
        "visits": len(visits),
        "stage_counts": report["stage_counts"],
        "cohort_counts": report["cohort_counts"],
        "validation_errors": report["validation"]["error_count"],
        "outputs": {key: str(value) for key, value in paths.items()},
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
