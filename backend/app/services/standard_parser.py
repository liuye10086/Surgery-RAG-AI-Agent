"""Structure-aware DOCX parsing for versioned medical standards."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument


_NUMBER = r"\d+(?:\.\d+)?"
_RANGE_RE = re.compile(rf"(?P<lower>{_NUMBER})\s*(?P<dash>[-~—～–至])\s*(?P<upper>{_NUMBER})(?P<unit>.*)$")
_UPPER_RE = re.compile(rf"(?P<op><|≤)\s*(?P<upper>{_NUMBER})(?P<unit>.*)$")
_LOWER_RE = re.compile(rf"(?P<op>>|≥)\s*(?P<lower>{_NUMBER})(?P<unit>.*)$")
_PERCENT_RE = re.compile(r"\s*%")
_SEX_RE = re.compile(r"(男性|女性)\s*([^；;]+)")


@dataclass(frozen=True)
class NumericExpression:
    lower: float | None = None
    upper: float | None = None
    lower_inclusive: bool = True
    upper_inclusive: bool = True
    unit: str = ""
    raw_text: str = ""


@dataclass(frozen=True)
class Segment:
    raw_text: str
    segment_type: str
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    column_index: int | None = None
    section_title: str | None = None
    source_metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class RuleCandidate:
    raw_text: str
    segment: Segment
    indicator_name: str | None = None
    target_state_type: str = "evidence"
    target_state_value: str | None = None
    rule_type: str = "qualitative_direction"
    machine_actionability: str = "evidence-only"
    evidence_type: str | None = None
    applicability: dict[str, Any] = field(default_factory=dict)
    numeric: NumericExpression | None = None
    interpretation: str | None = None


@dataclass(frozen=True)
class ParsedStandardDocument:
    path: Path
    parser_version: str
    paragraphs: list[Segment]
    tables: list[list[list[str]]]
    segments: list[Segment]
    rule_candidates: list[RuleCandidate]


def build_llm_candidate(segment_text: str, context: dict[str, Any] | None = None, adapter=None) -> dict[str, Any] | None:
    """Invoke an injected candidate adapter without coupling parsing to an LLM."""
    if adapter is None:
        return None
    try:
        result = adapter(segment_text, context or {})
    except Exception:
        return None
    return result if isinstance(result, dict) else None


def _clean(text: str) -> str:
    return " ".join(str(text or "").replace("\xa0", " ").split()).strip()


def _unit_from_tail(tail: str) -> str:
    tail = _clean(tail)
    if tail.startswith("%"):
        return "%" + tail[1:].strip()
    return tail


def parse_numeric_expression(text: str) -> NumericExpression | None:
    raw = _clean(text)
    if not raw:
        return None
    raw_without_fillers = re.sub(r"^(约|常见为|常作正常参考)\s*", "", raw)
    raw_without_fillers = raw_without_fillers.replace("\ufffd", "–")
    raw_without_fillers = re.sub(r"(?<=\d)\s*%?\s*[–—~～至]\s*(?=\d)", "–", raw_without_fillers)
    sex_parts = _SEX_RE.findall(raw_without_fillers)
    if sex_parts:
        raw_without_fillers = _clean(sex_parts[0][1])

    match = _RANGE_RE.search(raw_without_fillers)
    if match:
        return NumericExpression(
            lower=float(match.group("lower")),
            upper=float(match.group("upper")),
            lower_inclusive=True,
            upper_inclusive=True,
            unit=_unit_from_tail(match.group("unit")),
            raw_text=raw,
        )
    match = _UPPER_RE.search(raw_without_fillers)
    if match:
        return NumericExpression(
            upper=float(match.group("upper")),
            upper_inclusive=match.group("op") == "≤",
            unit=_unit_from_tail(match.group("unit")),
            raw_text=raw,
        )
    match = _LOWER_RE.search(raw_without_fillers)
    if match:
        return NumericExpression(
            lower=float(match.group("lower")),
            lower_inclusive=match.group("op") == "≥",
            unit=_unit_from_tail(match.group("unit")),
            raw_text=raw,
        )
    return None


def _section_for(paragraphs: list[str], index: int) -> str | None:
    section = None
    for text in paragraphs[: index + 1]:
        if re.match(r"^[一二三四五六七八九十]+、", text) or re.match(r"^\d+(?:\.\d+)?\s+", text):
            section = text
    return section


def _candidate_for_row(cells: list[str], segment: Segment) -> list[RuleCandidate]:
    if not cells:
        return []
    header = cells[0].lower()
    if any(token in header for token in ("field", "字段", "indicator_name", "metric_name", "判断方法")) or header in {"stage", "阶段"}:
        return []

    candidates: list[RuleCandidate] = []
    indicator = cells[0]
    joined = " | ".join(cells[1:])
    numeric = parse_numeric_expression(joined)
    state_type = "stage" if "stage" in indicator.lower() or indicator.lower().startswith("stage") else "evidence"
    if indicator.lower().startswith("stage") or (cells and any("biomarker_state" in cell.lower() for cell in cells)):
        state_type = "stage"
    if cells and any(cell.lower().startswith("stage ") for cell in cells):
        state_type = "stage"
    if len(cells) >= 5 and any(token in cells[0].lower() for token in ("lfc", "cap", "mri", "ctl/s", "histology")):
        state_type = "grade"
    if numeric:
        candidates.append(
            RuleCandidate(
                raw_text=segment.raw_text,
                segment=segment,
                indicator_name=indicator,
                target_state_type=state_type,
                target_state_value=cells[1] if len(cells) > 1 else None,
                rule_type="numeric_range",
                machine_actionability="calculable" if numeric.unit else "evidence-only",
                evidence_type="standard_table",
                numeric=numeric,
                interpretation=joined,
            )
        )
    else:
        candidates.append(
            RuleCandidate(
                raw_text=segment.raw_text,
                segment=segment,
                indicator_name=indicator,
                target_state_type=state_type,
                target_state_value=cells[1] if len(cells) > 1 else None,
                rule_type="classification" if state_type in {"stage", "grade"} else "qualitative_direction",
                machine_actionability="evidence-only",
                evidence_type="standard_table",
                interpretation=joined,
            )
        )
    return candidates


def parse_standard_docx(path: str | Path, *, parser_version: str) -> ParsedStandardDocument:
    source = Path(path)
    document = DocxDocument(str(source))
    paragraph_texts = [_clean(p.text) for p in document.paragraphs]
    paragraphs: list[Segment] = []
    segments: list[Segment] = []
    for index, text in enumerate(paragraph_texts):
        if not text:
            continue
        segment = Segment(
            raw_text=text,
            segment_type="paragraph" if not re.match(r"^\d+(?:\.\d+)?\s", text) else "rule_text",
            paragraph_index=index,
            section_title=_section_for(paragraph_texts, index),
        )
        paragraphs.append(segment)
        segments.append(segment)

    tables: list[list[list[str]]] = []
    candidates: list[RuleCandidate] = []
    for table_index, table in enumerate(document.tables):
        rows: list[list[str]] = []
        for row_index, row in enumerate(table.rows):
            cells = [_clean(cell.text) for cell in row.cells]
            rows.append(cells)
            raw_text = " | ".join(cells)
            if not raw_text:
                continue
            segment = Segment(
                raw_text=raw_text,
                segment_type="stage_row" if cells and cells[0].lower().startswith("stage") else "table_row",
                table_index=table_index,
                row_index=row_index,
                source_metadata={"column_count": len(cells)},
            )
            segments.append(segment)
            candidates.extend(_candidate_for_row(cells, segment))
        tables.append(rows)

    return ParsedStandardDocument(
        path=source,
        parser_version=parser_version,
        paragraphs=paragraphs,
        tables=tables,
        segments=segments,
        rule_candidates=candidates,
    )


def build_rule_candidates(parsed: ParsedStandardDocument) -> list[RuleCandidate]:
    return list(parsed.rule_candidates)
