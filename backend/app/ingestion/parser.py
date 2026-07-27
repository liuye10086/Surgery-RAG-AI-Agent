import io
import os
import re
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO, List, Union
from xml.etree import ElementTree as ET

from app.core.config import settings


# DOCX 命名空间
_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _clean_text(text: str) -> str:
    """清理提取出的文本：移除 NUL、统一换行、去首尾空白。"""
    if not text:
        return ""
    # PostgreSQL 不允许字符串中包含 NUL 字符
    text = text.replace("\x00", "")
    text = re.sub(r"\r\n|\r", "\n", text)
    return text.strip()


def _sanitize_docx_rels(src_path: str) -> io.BytesIO:
    """部分 WPS 生成的 docx 会在 relationships 里写 Target=\"NULL\"，
    python-docx 读取时会报 KeyError。此函数清理这类非法关系后，
    把修复后的文档直接写回内存（BytesIO），避免 Windows 下临时文件被占用。"""
    output = io.BytesIO()
    with zipfile.ZipFile(src_path, "r") as zin:
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".rels"):
                    try:
                        root = ET.fromstring(data)
                        changed = False
                        for rel in root.findall(f"{{{_RELS_NS}}}Relationship"):
                            target = rel.get("Target")
                            # WPS 有时会写 NULL、../NULL、空字符串等
                            target_name = Path(target).name if target else ""
                            if (
                                target is None
                                or target == ""
                                or target_name.upper() == "NULL"
                            ):
                                root.remove(rel)
                                changed = True
                        if changed:
                            # 保留 XML 声明
                            data = ET.tostring(
                                root,
                                encoding="UTF-8",
                                xml_declaration=True,
                            )
                    except Exception:
                        # XML 解析失败就原样写入
                        pass
                zout.writestr(item, data)
    output.seek(0)
    return output


def _extract_docx_text(file: Union[str, BinaryIO]) -> str:
    from docx import Document

    doc = Document(file)
    parts: List[str] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")

    return "\n".join(parts)


@dataclass
class ImageRef:
    """解析时提取的图片引用，blob 尚未落盘。"""
    blob: bytes
    ext: str        # png / jpg / jpeg ...
    page_number: int | None   # PDF 有页码，docx 为 None


@dataclass
class ExtractedPage:
    text: str
    page_number: int | None
    source_type: str
    images: list = None  # List[ImageRef]

    def __post_init__(self):
        if self.images is None:
            self.images = []


# PaddleOCR 实例懒加载，避免模块导入时就初始化
_ocr = None


def _get_ocr():
    global _ocr
    if _ocr is None:
        from paddleocr import PaddleOCR
        _ocr = PaddleOCR(
            use_angle_cls=True,
            lang=settings.PADDLEOCR_LANG,
            use_gpu=settings.PADDLEOCR_USE_GPU,
            show_log=False,
        )
    return _ocr


def _run_ocr(image_path: str) -> str:
    ocr = _get_ocr()
    result = ocr.ocr(image_path, cls=True)
    if not result or not result[0]:
        return ""
    lines = []
    for line in result[0]:
        if line and len(line) >= 2 and line[1] and len(line[1]) >= 1:
            lines.append(str(line[1][0]))
    return "\n".join(lines)


def parse_pdf(file_path: str) -> List[ExtractedPage]:
    import fitz  # pymupdf

    pages: List[ExtractedPage] = []
    doc = fitz.open(file_path)
    for page_idx in range(len(doc)):
        page = doc[page_idx]
        page_number = page_idx + 1
        text = _clean_text(page.get_text())

        # 提取页面内嵌图片
        page_images: list = []
        for img_info in page.get_images():
            xref = img_info[0]
            try:
                base = doc.extract_image(xref)
                ext = base.get("ext", "png")
                if ext == "jpeg":
                    ext = "jpg"
                page_images.append(ImageRef(
                    blob=base["image"],
                    ext=ext,
                    page_number=page_number,
                ))
            except Exception:
                pass  # 个别图片提取失败不阻塞整体解析

        if len(text) >= settings.PDF_OCR_MIN_TEXT_LENGTH:
            pages.append(
                ExtractedPage(
                    text=text,
                    page_number=page_number,
                    source_type="pdf_text",
                    images=page_images,
                )
            )
            continue

        # 低字数页走 OCR
        try:
            pix = page.get_pixmap(dpi=settings.PDF_OCR_DPI)
            fd, temp_path = tempfile.mkstemp(suffix=".png")
            os.close(fd)
            try:
                pix.save(temp_path)
                ocr_text = _clean_text(_run_ocr(temp_path))
            finally:
                os.unlink(temp_path)
        except Exception:
            ocr_text = ""

        pages.append(
            ExtractedPage(
                text=ocr_text or text,
                page_number=page_number,
                source_type="pdf_ocr" if ocr_text else "pdf_text",
                images=page_images,
            )
        )
    doc.close()
    return pages


def parse_docx(file_path: str) -> List[ExtractedPage]:
    # 先确认文件是 zip 格式（.docx 本质是 zip）
    if not zipfile.is_zipfile(file_path):
        raise ValueError(
            "无法解析该文件：它不是有效的 .docx 格式。"
            "可能是 .doc 老格式被强制改后缀，请用 Word/WPS 另存为 .docx 后重新上传。"
        )

    try:
        text = _clean_text(_extract_docx_text(file_path))
    except (KeyError, zipfile.BadZipFile) as e:
        # 常见 WPS 问题：relationships 里有 Target=NULL
        error_msg = str(e).lower()
        if "null" in error_msg or "rels" in error_msg or "item" in error_msg:
            try:
                sanitized = _sanitize_docx_rels(file_path)
                text = _clean_text(_extract_docx_text(sanitized))
            except Exception as inner:
                raise ValueError(
                    "该 .docx 文件结构异常，自动修复后仍无法解析。"
                    "建议用 Word/WPS 重新另存为 .docx 后上传。"
                ) from inner
        else:
            raise ValueError(
                f"解析 .docx 失败：{e}。建议重新另存为 .docx 后上传。"
            ) from e
    except Exception as e:
        raise ValueError(
            f"解析 .docx 失败：{e}。建议重新另存为 .docx 后上传。"
        ) from e

    images = _extract_docx_images(file_path)
    return [ExtractedPage(text=text, page_number=None, source_type="docx", images=images)]


def _extract_docx_images(file_path: str) -> list:
    """从 .docx 中按出现顺序提取内嵌图片，返回 ImageRef 列表。"""
    from docx import Document
    doc = Document(file_path)

    # 收集所有图片关系 rId → image_part
    image_parts: dict[str, object] = {}
    for r in doc.part.rels.values():
        if "image" in r.reltype:
            image_parts[r.rId] = r.target_part

    result = []
    NS_DRAWING = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS_BLIP = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    for para in doc.paragraphs:
        for run in para.runs:
            drawings = run._element.findall(f"{{{NS_DRAWING}}}drawing")
            for drawing in drawings:
                blips = drawing.findall(f".//{{{NS_BLIP}}}blip")
                for blip in blips:
                    embed = blip.get(f"{{{NS_R}}}embed")
                    if embed and embed in image_parts:
                        part = image_parts[embed]
                        ext = (part.content_type or "image/png").split("/")[-1]
                        if ext == "jpeg":
                            ext = "jpg"
                        result.append(ImageRef(
                            blob=part.blob,
                            ext=ext,
                            page_number=None,
                        ))
    return result


def parse_image(file_path: str) -> List[ExtractedPage]:
    text = _clean_text(_run_ocr(file_path))
    return [ExtractedPage(text=text, page_number=None, source_type="image_ocr")]


def parse_file(file_path: str) -> List[ExtractedPage]:
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    if ext == ".docx":
        return parse_docx(file_path)
    if ext in {".jpg", ".jpeg", ".png"}:
        return parse_image(file_path)
    if ext == ".doc":
        raise ValueError("暂不支持 .doc 格式，请转换为 .docx 后重新上传")
    raise ValueError(f"不支持的文件类型: {ext}")
